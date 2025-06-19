# core/letter_generator.py
"""
Module de génération de lettres et courriers juridiques.
"""
import os
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional, List, Tuple
from dataclasses import dataclass
import re

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

import streamlit as st
from core.llm.multi_llm_manager import MultiLLMManager
from core.vector_juridique import VectorJuridique


@dataclass
class LetterTemplate:
    """Représente un modèle de lettre."""
    name: str
    type: str
    description: str
    required_fields: List[str]
    optional_fields: List[str] = None
    tone: str = "formel"


class LetterGenerator:
    """Génère des lettres juridiques professionnelles."""
    
    def __init__(self, template_dir: str = "templates/lettres"):
        self.template_dir = Path(template_dir)
        self.template_dir.mkdir(parents=True, exist_ok=True)
        
        self.llm_manager = MultiLLMManager()
        self.vector_db = VectorJuridique()
        
        # Templates prédéfinis
        self.templates = {
            'mise_en_demeure': LetterTemplate(
                name="Mise en demeure",
                type="contentieux",
                description="Lettre de mise en demeure avec délai",
                required_fields=['destinataire', 'objet', 'faits', 'demande', 'delai'],
                optional_fields=['pieces_jointes', 'cc'],
                tone="assertif"
            ),
            'demande_pieces': LetterTemplate(
                name="Demande de pièces",
                type="procedure",
                description="Demande de communication de pièces",
                required_fields=['destinataire', 'dossier', 'pieces_demandees'],
                optional_fields=['delai', 'justification'],
                tone="formel"
            ),
            'courrier_client': LetterTemplate(
                name="Courrier client",
                type="client",
                description="Courrier d'information au client",
                required_fields=['destinataire', 'objet', 'contenu'],
                optional_fields=['rdv_propose', 'documents_joints'],
                tone="professionnel"
            ),
            'conclusions_depot': LetterTemplate(
                name="Bordereau de dépôt",
                type="juridiction",
                description="Bordereau de communication de conclusions",
                required_fields=['juridiction', 'numero_rg', 'parties', 'pieces'],
                optional_fields=['audience'],
                tone="formel"
            ),
            'custom': LetterTemplate(
                name="Lettre personnalisée",
                type="autre",
                description="Modèle libre",
                required_fields=['destinataire', 'objet', 'contenu'],
                tone="formel"
            )
        }
        
        # Styles de formatage
        self.styles = {
            'cabinet': {
                'name': 'STERU BARATTE AARPI',
                'address': '123 rue Example\n75001 PARIS',
                'phone': '01 23 45 67 89',
                'email': 'contact@steru-baratte.com',
                'color': RGBColor(30, 58, 138)  # Bleu foncé
            }
        }
    
    def get_available_templates(self) -> List[str]:
        """Retourne la liste des templates disponibles."""
        return list(self.templates.keys())
    
    def get_template_info(self, template_name: str) -> Optional[LetterTemplate]:
        """Retourne les informations d'un template."""
        return self.templates.get(template_name)
    
    async def generate_letter(
        self,
        template_name: str,
        fields: Dict[str, Any],
        use_ai: bool = True,
        ai_models: List[str] = None
    ) -> Tuple[Path, Dict[str, Any]]:
        """
        Génère une lettre basée sur un template.
        
        Args:
            template_name: Nom du template à utiliser
            fields: Champs à remplir
            use_ai: Utiliser l'IA pour enrichir le contenu
            ai_models: Modèles IA à utiliser
        
        Returns:
            (chemin_fichier, metadata)
        """
        template = self.templates.get(template_name)
        if not template:
            raise ValueError(f"Template inconnu : {template_name}")
        
        # Vérifier les champs requis
        missing_fields = [f for f in template.required_fields if f not in fields]
        if missing_fields:
            raise ValueError(f"Champs manquants : {', '.join(missing_fields)}")
        
        # Enrichir avec l'IA si demandé
        if use_ai:
            fields = await self._enrich_with_ai(template, fields, ai_models or ['GPT-4o'])
        
        # Créer le document
        doc = Document()
        self._setup_document_styles(doc)
        
        # En-tête cabinet
        self._add_header(doc)
        
        # Date et lieu
        self._add_date_location(doc, fields.get('lieu', 'Paris'))
        
        # Destinataire
        self._add_recipient(doc, fields['destinataire'])
        
        # Objet
        if 'objet' in fields:
            self._add_subject(doc, fields['objet'])
        
        # Corps de la lettre
        self._add_letter_body(doc, template, fields)
        
        # Signature
        self._add_signature(doc, fields.get('signataire', 'Me Edouard STERU'))
        
        # Pièces jointes
        if 'pieces_jointes' in fields:
            self._add_attachments(doc, fields['pieces_jointes'])
        
        # Copie conforme
        if 'cc' in fields:
            self._add_cc(doc, fields['cc'])
        
        # Sauvegarder
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{template_name}_{timestamp}.docx"
        filepath = self.template_dir / filename
        doc.save(str(filepath))
        
        # Métadonnées
        metadata = {
            'template': template_name,
            'generated_at': datetime.now().isoformat(),
            'filepath': str(filepath),
            'fields_used': list(fields.keys()),
            'ai_enhanced': use_ai,
            'file_size': filepath.stat().st_size
        }
        
        return filepath, metadata
    
    async def _enrich_with_ai(
        self,
        template: LetterTemplate,
        fields: Dict[str, Any],
        ai_models: List[str]
    ) -> Dict[str, Any]:
        """Enrichit les champs avec l'IA."""
        # Construire le prompt
        prompt = f"""En tant qu'avocat spécialisé en droit pénal des affaires, rédigez une {template.name} professionnelle.

Informations fournies :
{self._format_fields_for_prompt(fields)}

Ton souhaité : {template.tone}

Consignes :
1. Utiliser un langage juridique précis
2. Structurer clairement les arguments
3. Respecter le ton {template.tone}
4. Inclure les références légales pertinentes si applicable

Format de réponse attendu en JSON :
{{
    "contenu_enrichi": "...",
    "formules_juridiques": ["...", "..."],
    "references_legales": ["...", "..."]
}}"""

        # Interroger l'IA
        response = await self.llm_manager.query_multiple(
            prompt=prompt,
            context="",
            selected_models=ai_models
        )
        
        # Parser et enrichir
        try:
            import json
            # Prendre la première réponse valide
            for model, resp in response['responses'].items():
                if not resp.get('error'):
                    content = resp.get('content', '')
                    # Extraire le JSON
                    json_match = re.search(r'\{.*\}', content, re.DOTALL)
                    if json_match:
                        ai_data = json.loads(json_match.group())
                        
                        # Enrichir les champs
                        if 'contenu' in fields and ai_data.get('contenu_enrichi'):
                            fields['contenu'] = ai_data['contenu_enrichi']
                        
                        fields['formules_juridiques'] = ai_data.get('formules_juridiques', [])
                        fields['references_legales'] = ai_data.get('references_legales', [])
                        
                        break
        except:
            # En cas d'erreur, garder les champs originaux
            pass
        
        return fields
    
    def _format_fields_for_prompt(self, fields: Dict[str, Any]) -> str:
        """Formate les champs pour le prompt IA."""
        lines = []
        for key, value in fields.items():
            if isinstance(value, list):
                value = ', '.join(str(v) for v in value)
            lines.append(f"- {key}: {value}")
        return '\n'.join(lines)
    
    def _setup_document_styles(self, doc: Document):
        """Configure les styles du document."""
        # Style pour l'en-tête
        styles = doc.styles
        
        # Créer ou modifier le style de titre
        if 'CustomTitle' not in styles:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        else:
            title_style = styles['CustomTitle']
        
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(14)
        title_style.font.bold = True
        title_style.font.color.rgb = self.styles['cabinet']['color']
    
    def _add_header(self, doc: Document):
        """Ajoute l'en-tête du cabinet."""
        # Logo et informations du cabinet
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Nom du cabinet
        run = header.add_run(self.styles['cabinet']['name'])
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = self.styles['cabinet']['color']
        
        header.add_run('\n' + self.styles['cabinet']['address'])
        header.add_run('\n' + self.styles['cabinet']['phone'])
        header.add_run('\n' + self.styles['cabinet']['email'])
        
        # Ligne de séparation
        doc.add_paragraph('_' * 80)
    
    def _add_date_location(self, doc: Document, location: str):
        """Ajoute la date et le lieu."""
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_text = f"{location}, le {datetime.now().strftime('%d %B %Y')}"
        date_para.add_run(date_text)
    
    def _add_recipient(self, doc: Document, recipient: str):
        """Ajoute le destinataire."""
        doc.add_paragraph()  # Espace
        recipient_para = doc.add_paragraph(recipient)
        recipient_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _add_subject(self, doc: Document, subject: str):
        """Ajoute l'objet de la lettre."""
        doc.add_paragraph()  # Espace
        subject_para = doc.add_paragraph()
        subject_para.add_run('Objet : ').bold = True
        subject_para.add_run(subject)
    
    def _add_letter_body(self, doc: Document, template: LetterTemplate, fields: Dict[str, Any]):
        """Ajoute le corps de la lettre."""
        doc.add_paragraph()  # Espace
        
        # Formule de politesse d'ouverture
        if template.tone == "formel":
            doc.add_paragraph("Madame, Monsieur,")
        elif template.tone == "assertif":
            doc.add_paragraph("Maître,")
        else:
            doc.add_paragraph("Cher(e) client(e),")
        
        doc.add_paragraph()  # Espace
        
        # Contenu principal
        content = fields.get('contenu', '')
        
        # Si c'est une mise en demeure, structurer différemment
        if template.type == "contentieux":
            # Rappel des faits
            if 'faits' in fields:
                faits_para = doc.add_paragraph()
                faits_para.add_run("I. RAPPEL DES FAITS\n").bold = True
                doc.add_paragraph(fields['faits'])
            
            # Demande
            if 'demande' in fields:
                demande_para = doc.add_paragraph()
                demande_para.add_run("II. DEMANDE\n").bold = True
                doc.add_paragraph(fields['demande'])
            
            # Délai
            if 'delai' in fields:
                doc.add_paragraph()
                delai_text = f"Vous disposez d'un délai de {fields['delai']} pour vous conformer à cette mise en demeure."
                doc.add_paragraph(delai_text)
        else:
            # Contenu libre
            paragraphs = content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    doc.add_paragraph(para_text)
        
        # Références légales si présentes
        if 'references_legales' in fields and fields['references_legales']:
            doc.add_paragraph()
            ref_para = doc.add_paragraph()
            ref_para.add_run("Références légales : ").bold = True
            for ref in fields['references_legales']:
                doc.add_paragraph(f"- {ref}", style='List Bullet')
        
        # Formule de conclusion
        doc.add_paragraph()
        if template.tone == "formel":
            doc.add_paragraph("Je vous prie d'agréer, Madame, Monsieur, l'expression de mes salutations distinguées.")
        elif template.tone == "assertif":
            doc.add_paragraph("Dans l'attente de votre réponse, je vous prie d'agréer, Maître, l'expression de mes salutations confraternelles.")
        else:
            doc.add_paragraph("Restant à votre disposition pour tout renseignement complémentaire.")
    
    def _add_signature(self, doc: Document, signatory: str):
        """Ajoute la signature."""
        doc.add_paragraph()  # Espace
        doc.add_paragraph()  # Espace
        
        sig_para = doc.add_paragraph()
        sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig_para.add_run(signatory)
    
    def _add_attachments(self, doc: Document, attachments: List[str]):
        """Ajoute la liste des pièces jointes."""
        doc.add_paragraph()  # Espace
        doc.add_paragraph()  # Espace
        
        pj_para = doc.add_paragraph()
        pj_para.add_run("P.J. : ").bold = True
        
        for attachment in attachments:
            doc.add_paragraph(f"- {attachment}", style='List Bullet')
    
    def _add_cc(self, doc: Document, cc_list: List[str]):
        """Ajoute les copies conformes."""
        doc.add_paragraph()  # Espace
        
        cc_para = doc.add_paragraph()
        cc_para.add_run("Copie : ").bold = True
        
        for cc in cc_list:
            doc.add_paragraph(f"- {cc}", style='List Bullet')
    
    def preview_letter(self, filepath: Path) -> str:
        """Génère un aperçu texte de la lettre."""
        doc = Document(str(filepath))
        
        preview_lines = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                preview_lines.append(paragraph.text)
        
        return '\n\n'.join(preview_lines)


# Fonction helper pour Streamlit
def generate_letter(
    letter_type: str,
    context: Dict[str, Any]
) -> Optional[Path]:
    """
    Fonction wrapper pour générer une lettre depuis Streamlit.
    
    Args:
        letter_type: Type de lettre
        context: Contexte avec les champs
    
    Returns:
        Chemin du fichier généré ou None
    """
    generator = LetterGenerator()
    
    try:
        # Adapter le contexte aux champs attendus
        fields = {
            'destinataire': context.get('destinataire', 'À définir'),
            'objet': context.get('objet', f'Lettre {letter_type}'),
            'contenu': context.get('contenu', ''),
            'lieu': context.get('lieu', 'Paris'),
            'signataire': context.get('signataire', 'Me Edouard STERU')
        }
        
        # Ajouter les champs spécifiques selon le type
        if letter_type == 'mise_en_demeure':
            fields.update({
                'faits': context.get('faits', ''),
                'demande': context.get('demande', ''),
                'delai': context.get('delai', '8 jours')
            })
        
        # Générer de manière synchrone pour Streamlit
        import asyncio
        filepath, metadata = asyncio.run(
            generator.generate_letter(
                template_name=letter_type,
                fields=fields,
                use_ai=context.get('use_ai', False)
            )
        )
        
        return filepath
        
    except Exception as e:
        st.error(f"Erreur lors de la génération : {e}")
        return None


# Export
__all__ = ['LetterGenerator', 'generate_letter', 'LetterTemplate']
