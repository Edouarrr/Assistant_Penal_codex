"""
Module de génération d'actes juridiques avec gestion avancée des pièces.
"""
import re
import os
import json
import shutil
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
from dataclasses import dataclass, field
import hashlib

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from core.llm.multi_llm_manager import MultiLLMManager
from core.vector_juridique import VectorJuridique
from core.search.intelligent_search import IntelligentSearch


@dataclass
class Piece:
    """Représente une pièce juridique."""
    numero: int
    titre: str
    fichier_source: str
    fichier_communique: Optional[str] = None
    type: str = "versee"  # versee, adverse, penale
    pages: List[int] = field(default_factory=list)
    hash: Optional[str] = None
    premiere_citation: bool = True


@dataclass
class Citation:
    """Représente une citation dans le document."""
    texte: str
    piece: Piece
    page: Optional[int] = None
    position: int = 0  # Position dans le document


class LegalDocumentGenerator:
    """Génère des documents juridiques avec gestion des pièces et citations."""
    
    def __init__(self, template_dir: str = "templates/actes"):
        self.template_dir = Path(template_dir)
        self.template_dir.mkdir(parents=True, exist_ok=True)
        
        self.pieces_dir = Path("pieces_communiquees")
        self.pieces_dir.mkdir(parents=True, exist_ok=True)
        
        self.llm_manager = MultiLLMManager()
        self.vector_db = VectorJuridique()
        self.search_engine = IntelligentSearch()
        
        # Registre des pièces
        self.pieces_registry: Dict[str, Piece] = {}
        self.citations: List[Citation] = []
        
        # Styles de formatage
        self.styles = {
            'piece_number': {'bold': True, 'highlight': 'yellow'},
            'footnote': {'size': 9, 'italic': True},
            'heading1': {'size': 16, 'bold': True},
            'heading2': {'size': 14, 'bold': True},
            'normal': {'size': 11}
        }
    
    def generate_act(
        self,
        act_type: str,
        context: Dict[str, Any],
        reference_docs: List[str],
        key_points: List[str],
        models: List[str] = None,
        options: Dict[str, Any] = None
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Génère un acte juridique complet.
        
        Args:
            act_type: Type d'acte (conclusions, plainte, etc.)
            context: Contexte (juridiction, n° dossier, etc.)
            reference_docs: Documents de référence
            key_points: Points clés à développer
            models: Modèles LLM à utiliser
            options: Options de génération
        
        Returns:
            Tuple (chemin du fichier généré, métadonnées)
        """
        if models is None:
            models = ['GPT-4o', 'Claude Opus 4']
        
        if options is None:
            options = {}
        
        # 1. Analyser les documents de référence
        analysis = self._analyze_reference_documents(reference_docs)
        
        # 2. Construire le prompt enrichi
        prompt = self._build_generation_prompt(
            act_type, context, key_points, analysis, options
        )
        
        # 3. Interroger les LLM
        llm_responses = self._query_llms(prompt, analysis['context'], models)
        
        # 4. Fusionner et structurer le contenu
        structured_content = self._structure_content(
            act_type, llm_responses, analysis
        )
        
        # 5. Extraire et enregistrer les pièces
        self._extract_and_register_pieces(structured_content, reference_docs)
        
        # 6. Générer le document Word
        doc_path = self._generate_word_document(
            act_type, context, structured_content, options
        )
        
        # 7. Copier et renommer les pièces communiquées
        pieces_info = self._process_communicated_pieces()
        
        # 8. Générer les métadonnées
        metadata = {
            'act_type': act_type,
            'generation_date': datetime.now().isoformat(),
            'context': context,
            'pieces_count': len(self.pieces_registry),
            'pieces_communicated': pieces_info,
            'models_used': models,
            'word_count': self._count_words(structured_content),
            'pages_estimate': self._estimate_pages(structured_content)
        }
        
        return doc_path, metadata
    
    def _analyze_reference_documents(
        self, 
        reference_docs: List[str]
    ) -> Dict[str, Any]:
        """Analyse les documents de référence."""
        analysis = {
            'summaries': [],
            'key_facts': [],
            'contradictions': [],
            'citations': [],
            'context': ""
        }
        
        # Rechercher et analyser chaque document
        for doc_name in reference_docs:
            # Recherche dans la base vectorielle
            results = self.vector_db.search(
                query=f"document:{doc_name}",
                k=20,
                filter_dict={'file_name': {'$contains': doc_name}}
            )
            
            if results:
                # Extraire les informations clés
                doc_summary = self._summarize_document(results)
                analysis['summaries'].append({
                    'document': doc_name,
                    'summary': doc_summary,
                    'chunks': len(results)
                })
                
                # Extraire les faits importants
                facts = self._extract_key_facts(results)
                analysis['key_facts'].extend(facts)
                
                # Préparer le contexte pour les LLM
                for result in results[:5]:  # Top 5 chunks
                    analysis['context'] += f"\n--- {doc_name} ---\n"
                    analysis['context'] += result['content']
                    analysis['context'] += "\n---\n"
        
        return analysis
    
    def _build_generation_prompt(
        self,
        act_type: str,
        context: Dict[str, Any],
        key_points: List[str],
        analysis: Dict[str, Any],
        options: Dict[str, Any]
    ) -> str:
        """Construit le prompt pour la génération."""
        prompt_parts = []
        
        # Instructions de base
        prompt_parts.append(
            f"Tu es un avocat expert en droit pénal des affaires. "
            f"Tu dois rédiger {act_type} pour {context.get('jurisdiction', 'le tribunal')}."
        )
        
        # Style et ton
        tone_map = {
            'Très formel': "utilise un style très soutenu et protocolaire",
            'Formel': "utilise un style formel et professionnel",
            'Assertif': "utilise un style assertif et convaincant",
            'Combatif': "utilise un style combatif et offensif"
        }
        
        tone = options.get('tone', 'Formel')
        prompt_parts.append(f"Ton : {tone_map.get(tone, tone_map['Formel'])}")
        
        # Contexte de l'affaire
        prompt_parts.append("\nContexte de l'affaire :")
        prompt_parts.append(f"- N° de dossier : {context.get('case_number', 'XXX')}")
        prompt_parts.append(f"- Magistrat : {context.get('judge_name', 'XXX')}")
        
        # Points clés à développer
        prompt_parts.append("\nPoints clés à développer impérativement :")
        for point in key_points:
            prompt_parts.append(f"- {point}")
        
        # Structure attendue
        prompt_parts.append(f"\nStructure attendue pour {act_type} :")
        
        structure = self._get_document_structure(act_type)
        for section in structure:
            prompt_parts.append(f"- {section}")
        
        # Instructions spécifiques
        prompt_parts.append("\nInstructions importantes :")
        prompt_parts.append("- Cite précisément les pièces avec leur numéro")
        prompt_parts.append("- Utilise des arguments juridiques solides")
        prompt_parts.append("- Structure le document avec des titres clairs")
        prompt_parts.append("- Inclus les articles de loi pertinents")
        
        if options.get('include_jurisprudence'):
            prompt_parts.append("- Cite la jurisprudence récente pertinente")
        
        # Résumé des documents analysés
        if analysis['summaries']:
            prompt_parts.append("\nRésumé des documents de référence :")
            for summary in analysis['summaries']:
                prompt_parts.append(
                    f"- {summary['document']} : {summary['summary'][:200]}..."
                )
        
        return "\n".join(prompt_parts)
    
    def _get_document_structure(self, act_type: str) -> List[str]:
        """Retourne la structure type d'un document."""
        structures = {
            "Conclusions (défense)": [
                "EN-TÊTE (Juridiction, parties, n° RG)",
                "RAPPEL DES FAITS ET DE LA PROCÉDURE",
                "DISCUSSION",
                "I. SUR LA RÉGULARITÉ DE LA PROCÉDURE",
                "II. SUR LE FOND",
                "A. Sur l'absence d'éléments constitutifs",
                "B. Sur les moyens de défense",
                "PAR CES MOTIFS",
                "LISTE DES PIÈCES COMMUNIQUÉES"
            ],
            "Plainte avec constitution de partie civile": [
                "EN-TÊTE (Doyen des juges d'instruction)",
                "EXPOSÉ DES FAITS",
                "QUALIFICATION JURIDIQUE",
                "PRÉJUDICE SUBI",
                "CONSTITUTION DE PARTIE CIVILE",
                "DEMANDE D'ACTES",
                "LISTE DES PIÈCES JOINTES"
            ],
            "QPC": [
                "EN-TÊTE",
                "OBJET : Question prioritaire de constitutionnalité",
                "RAPPEL DES FAITS ET DE LA PROCÉDURE",
                "SUR LA RECEVABILITÉ DE LA QPC",
                "SUR LE CARACTÈRE SÉRIEUX DE LA QPC",
                "I. Sur la méconnaissance du principe...",
                "II. Sur l'atteinte aux droits et libertés",
                "PAR CES MOTIFS"
            ]
        }
        
        return structures.get(act_type, ["INTRODUCTION", "DÉVELOPPEMENT", "CONCLUSION"])
    
    def _query_llms(
        self,
        prompt: str,
        context: str,
        models: List[str]
    ) -> Dict[str, Any]:
        """Interroge les LLM pour générer le contenu."""
        # Pour la démo, simulation
        # En production : utiliser self.llm_manager.query_multiple()
        
        responses = {}
        
        for model in models:
            responses[model] = {
                'content': self._generate_mock_content(prompt, model),
                'confidence': 0.9,
                'tokens_used': 2500
            }
        
        return responses
    
    def _generate_mock_content(self, prompt: str, model: str) -> str:
        """Génère un contenu simulé pour la démo."""
        return f"""
# CONCLUSIONS EN DÉFENSE

## POUR : Monsieur Jean MARTIN

## CONTRE : Ministère Public

---

## RAPPEL DES FAITS ET DE LA PROCÉDURE

Par réquisitoire introductif en date du 15 janvier 2024, le Procureur de la République a requis l'ouverture d'une information judiciaire des chefs d'escroquerie et abus de confiance.

Il résulte de la procédure que mon client aurait, selon l'accusation, détourné des fonds appartenant à la société XYZ (Pièce n°1 - PV d'audition du 20/01/2024).

## DISCUSSION

### I. SUR LA RÉGULARITÉ DE LA PROCÉDURE

#### A. Sur la nullité du procès-verbal d'audition

Il convient de relever que le procès-verbal d'audition de mon client en date du 20 janvier 2024 (Pièce n°2) est entaché de nullité.

En effet, il n'a pas été notifié à mon client son droit de se taire, en violation de l'article 63-1 du Code de procédure pénale.

### II. SUR LE FOND

#### A. Sur l'absence d'éléments constitutifs de l'escroquerie

L'escroquerie suppose la réunion de plusieurs éléments constitutifs qui font défaut en l'espèce.

Selon l'article 313-1 du Code pénal, l'escroquerie est "le fait, soit par l'usage d'un faux nom ou d'une fausse qualité, soit par l'abus d'une qualité vraie, soit par l'emploi de manœuvres frauduleuses, de tromper une personne physique ou morale".

Or, il ressort du rapport d'expertise comptable (Pièce n°3) que les mouvements de fonds incriminés correspondent à des opérations régulières.

## PAR CES MOTIFS

Il est demandé au Tribunal de :

- CONSTATER la nullité du procès-verbal d'audition du 20 janvier 2024
- DIRE ET JUGER que les éléments constitutifs de l'escroquerie ne sont pas réunis
- PRONONCER la relaxe de Monsieur Jean MARTIN

---

Généré par {model} pour démonstration
"""
    
    def _structure_content(
        self,
        act_type: str,
        llm_responses: Dict[str, Any],
        analysis: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Structure et fusionne le contenu des LLM."""
        # Fusion des réponses
        fused_content = self._fuse_llm_responses(llm_responses)
        
        # Structuration par sections
        sections = self._parse_sections(fused_content)
        
        # Enrichissement avec les analyses
        enriched_sections = self._enrich_sections(sections, analysis)
        
        # Validation de la structure
        validated_structure = self._validate_structure(act_type, enriched_sections)
        
        return validated_structure
    
    def _fuse_llm_responses(self, responses: Dict[str, Any]) -> str:
        """Fusionne les réponses des différents LLM."""
        # Pour la démo, on prend la première réponse
        # En production, fusion intelligente
        
        first_model = list(responses.keys())[0]
        return responses[first_model]['content']
    
    def _parse_sections(self, content: str) -> Dict[str, str]:
        """Parse le contenu en sections."""
        sections = {}
        current_section = "INTRODUCTION"
        current_content = []
        
        # Patterns pour détecter les sections
        section_patterns = [
            r'^#{1,3}\s+(.+)$',  # Markdown headers
            r'^[IVX]+\.\s+(.+)$',  # Roman numerals
            r'^[A-Z][A-Z\s]+:?\s*$',  # ALL CAPS headers
        ]
        
        lines = content.split('\n')
        
        for line in lines:
            is_header = False
            
            for pattern in section_patterns:
                match = re.match(pattern, line.strip())
                if match:
                    # Sauvegarder la section précédente
                    if current_content:
                        sections[current_section] = '\n'.join(current_content)
                    
                    # Nouvelle section
                    current_section = match.group(1) if match.lastindex else line.strip()
                    current_content = []
                    is_header = True
                    break
            
            if not is_header and line.strip():
                current_content.append(line)
        
        # Dernière section
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    
    def _enrich_sections(
        self,
        sections: Dict[str, str],
        analysis: Dict[str, Any]
    ) -> Dict[str, str]:
        """Enrichit les sections avec les analyses."""
        enriched = sections.copy()
        
        # Ajouter les faits clés si manquants
        if analysis['key_facts'] and 'FAITS' in enriched:
            facts_section = enriched['FAITS']
            for fact in analysis['key_facts'][:5]:  # Top 5 facts
                if fact not in facts_section:
                    facts_section += f"\n\n{fact}"
            enriched['FAITS'] = facts_section
        
        return enriched
    
    def _validate_structure(
        self,
        act_type: str,
        sections: Dict[str, str]
    ) -> Dict[str, Any]:
        """Valide et complète la structure du document."""
        required_structure = self._get_document_structure(act_type)
        
        validated = {
            'type': act_type,
            'sections': []
        }
        
        # Vérifier chaque section requise
        for required_section in required_structure:
            found = False
            
            for section_title, content in sections.items():
                if self._section_matches(required_section, section_title):
                    validated['sections'].append({
                        'title': required_section,
                        'content': content
                    })
                    found = True
                    break
            
            if not found:
                # Ajouter une section vide si manquante
                validated['sections'].append({
                    'title': required_section,
                    'content': f"[{required_section} - À compléter]"
                })
        
        return validated
    
    def _section_matches(self, required: str, actual: str) -> bool:
        """Vérifie si une section correspond à celle requise."""
        # Normaliser pour la comparaison
        required_normalized = required.lower().strip()
        actual_normalized = actual.lower().strip()
        
        # Correspondance exacte
        if required_normalized == actual_normalized:
            return True
        
        # Correspondance partielle
        if required_normalized in actual_normalized or actual_normalized in required_normalized:
            return True
        
        # Mots clés
        keywords_map = {
            'faits': ['fait', 'rappel', 'exposé', 'circonstance'],
            'discussion': ['discussion', 'moyens', 'arguments'],
            'motifs': ['motif', 'demande', 'conclusion', 'par ces motifs'],
            'pièces': ['pièce', 'annexe', 'justificatif']
        }
        
        for key, keywords in keywords_map.items():
            if key in required_normalized:
                return any(kw in actual_normalized for kw in keywords)
        
        return False
    
    def _extract_and_register_pieces(
        self,
        structured_content: Dict[str, Any],
        reference_docs: List[str]
    ) -> None:
        """Extrait et enregistre les pièces citées."""
        piece_counter = 1
        
        # Pattern pour détecter les citations de pièces
        piece_patterns = [
            r'[Pp]ièce\s*n°\s*(\d+)',
            r'\(Pièce\s*n°\s*(\d+)[^)]*\)',
            r'pièce\s+(\d+)',
        ]
        
        # Parcourir toutes les sections
        for section in structured_content['sections']:
            content = section['content']
            
            # Chercher les citations de pièces
            for pattern in piece_patterns:
                matches = re.finditer(pattern, content)
                
                for match in matches:
                    piece_num = int(match.group(1))
                    
                    # Extraire le contexte pour identifier la pièce
                    start = max(0, match.start() - 50)
                    end = min(len(content), match.end() + 100)
                    context = content[start:end]
                    
                    # Identifier la pièce
                    piece_title = self._extract_piece_title(context, piece_num)
                    
                    # Chercher le fichier source
                    source_file = self._find_source_file(piece_title, reference_docs)
                    
                    # Enregistrer la pièce
                    if piece_num not in self.pieces_registry:
                        piece = Piece(
                            numero=piece_num,
                            titre=piece_title,
                            fichier_source=source_file or "inconnu",
                            type="versee"
                        )
                        self.pieces_registry[piece_num] = piece
                    
                    # Enregistrer la citation
                    citation = Citation(
                        texte=match.group(0),
                        piece=self.pieces_registry[piece_num],
                        position=match.start()
                    )
                    self.citations.append(citation)
    
    def _extract_piece_title(self, context: str, piece_num: int) -> str:
        """Extrait le titre d'une pièce depuis son contexte."""
        # Patterns pour extraire le titre
        title_patterns = [
            rf'Pièce\s*n°\s*{piece_num}\s*[-–:]\s*([^,\.\n]+)',
            rf'Pièce\s*n°\s*{piece_num}\s*\(([^)]+)\)',
            rf'([^,\.\n]+)\s*\(Pièce\s*n°\s*{piece_num}\)',
        ]
        
        for pattern in title_patterns:
            match = re.search(pattern, context, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        # Titre par défaut
        return f"Document {piece_num}"
    
    def _find_source_file(
        self,
        piece_title: str,
        reference_docs: List[str]
    ) -> Optional[str]:
        """Trouve le fichier source d'une pièce."""
        # Normaliser le titre
        title_normalized = piece_title.lower()
        
        # Chercher dans les documents de référence
        for doc in reference_docs:
            doc_normalized = doc.lower()
            
            # Correspondance exacte ou partielle
            if (title_normalized in doc_normalized or 
                doc_normalized in title_normalized or
                self._similarity_score(title_normalized, doc_normalized) > 0.7):
                return doc
        
        # Chercher dans la base vectorielle
        results = self.vector_db.search(
            query=piece_title,
            k=1,
            filter_dict={'document_type': {'$in': ['audition', 'expertise', 'judiciaire']}}
        )
        
        if results:
            return results[0]['metadata'].get('file_name')
        
        return None
    
    def _similarity_score(self, str1: str, str2: str) -> float:
        """Calcule un score de similarité entre deux chaînes."""
        # Simplification : ratio de mots communs
        words1 = set(str1.split())
        words2 = set(str2.split())
        
        if not words1 or not words2:
            return 0.0
        
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        
        return len(intersection) / len(union)
    
    def _generate_word_document(
        self,
        act_type: str,
        context: Dict[str, Any],
        structured_content: Dict[str, Any],
        options: Dict[str, Any]
    ) -> str:
        """Génère le document Word final."""
        # Créer un nouveau document
        doc = Document()
        
        # Configurer les styles
        self._setup_document_styles(doc)
        
        # En-tête
        self._add_header(doc, act_type, context)
        
        # Corps du document
        for section in structured_content['sections']:
            self._add_section(doc, section)
        
        # Liste des pièces si demandée
        if options.get('include_pieces', True):
            self._add_pieces_list(doc)
        
        # Métadonnées
        self._add_document_properties(doc, act_type, context)
        
        # Sauvegarder
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{act_type.replace(' ', '_')}_{timestamp}.docx"
        filepath = Path("generated_documents") / filename
        filepath.parent.mkdir(exist_ok=True)
        
        doc.save(str(filepath))
        
        return str(filepath)
    
    def _setup_document_styles(self, doc: Document) -> None:
        """Configure les styles du document."""
        styles = doc.styles
        
        # Style pour les numéros de pièces
        if 'PieceNumber' not in styles:
            piece_style = styles.add_style('PieceNumber', WD_STYLE_TYPE.CHARACTER)
            piece_style.font.bold = True
            piece_style.font.highlight_color = 7  # Jaune
        
        # Style pour les notes de bas de page
        if 'PieceFootnote' not in styles:
            footnote_style = styles.add_style('PieceFootnote', WD_STYLE_TYPE.CHARACTER)
            footnote_style.font.size = Pt(9)
            footnote_style.font.italic = True
    
    def _add_header(self, doc: Document, act_type: str, context: Dict[str, Any]) -> None:
        """Ajoute l'en-tête du document."""
        # Titre principal
        title = doc.add_paragraph(act_type.upper())
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(16)
        title.runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Juridiction
        if context.get('jurisdiction'):
            juridiction = doc.add_paragraph(f"DEVANT {context['jurisdiction'].upper()}")
            juridiction.alignment = WD_ALIGN_PARAGRAPH.CENTER
            juridiction.runs[0].font.size = Pt(14)
        
        # Numéro de dossier
        if context.get('case_number'):
            case_para = doc.add_paragraph(f"N° RG : {context['case_number']}")
            case_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ligne de séparation
        doc.add_paragraph('_' * 50)
        doc.add_paragraph()
    
    def _add_section(self, doc: Document, section: Dict[str, str]) -> None:
        """Ajoute une section au document."""
        # Titre de section
        if not section['title'].startswith('['):  # Pas une section placeholder
            heading = doc.add_heading(section['title'], level=2)
            heading.runs[0].font.size = Pt(14)
            heading.runs[0].font.bold = True
        
        # Contenu avec formatage des pièces
        content = section['content']
        
        # Traiter les paragraphes
        paragraphs = content.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph()
                
                # Traiter les citations de pièces
                self._format_paragraph_with_pieces(para, para_text)
    
    def _format_paragraph_with_pieces(self, paragraph, text: str) -> None:
        """Formate un paragraphe avec les citations de pièces."""
        # Pattern pour les pièces
        piece_pattern = r'(\()?[Pp]ièce\s*n°\s*(\d+)([^)]*\))?'
        
        last_end = 0
        
        for match in re.finditer(piece_pattern, text):
            # Ajouter le texte avant la pièce
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])
            
            # Formater la citation de pièce
            piece_num = int(match.group(2))
            piece_text = match.group(0)
            
            # Ajouter avec style
            piece_run = paragraph.add_run(piece_text)
            piece_run.font.bold = True
            
            # Surligner si c'est une pièce versée
            if piece_num in self.pieces_registry:
                piece = self.pieces_registry[piece_num]
                if piece.type == "versee":
                    # Ajouter le surlignage jaune
                    piece_run.font.highlight_color = 7
                
                # Note de bas de page à la première citation
                if piece.premiere_citation:
                    self._add_footnote(paragraph, piece.titre)
                    piece.premiere_citation = False
            
            last_end = match.end()
        
        # Ajouter le reste du texte
        if last_end < len(text):
            paragraph.add_run(text[last_end:])
    
    def _add_footnote(self, paragraph, footnote_text: str) -> None:
        """Ajoute une note de bas de page."""
        # Simplification : ajouter entre parenthèses
        # En production, utiliser les vraies footnotes Word
        footnote_run = paragraph.add_run(f" ({footnote_text})")
        footnote_run.font.size = Pt(9)
        footnote_run.font.italic = True
    
    def _add_pieces_list(self, doc: Document) -> None:
        """Ajoute la liste des pièces communiquées."""
        doc.add_page_break()
        
        # Titre
        title = doc.add_heading("LISTE DES PIÈCES COMMUNIQUÉES", level=1)
        title.runs[0].font.size = Pt(16)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Trier les pièces par numéro
        sorted_pieces = sorted(
            [(num, piece) for num, piece in self.pieces_registry.items() 
             if piece.type == "versee"],
            key=lambda x: x[0]
        )
        
        # Ajouter chaque pièce
        for num, piece in sorted_pieces:
            piece_para = doc.add_paragraph()
            
            # Numéro en gras
            num_run = piece_para.add_run(f"Pièce n°{num:03d} : ")
            num_run.font.bold = True
            
            # Titre de la pièce
            piece_para.add_run(piece.titre)
            
            # Ajouter les pages si disponibles
            if piece.pages:
                piece_para.add_run(f" (pages {', '.join(map(str, piece.pages))})")
    
    def _add_document_properties(
        self,
        doc: Document,
        act_type: str,
        context: Dict[str, Any]
    ) -> None:
        """Ajoute les propriétés du document."""
        core_properties = doc.core_properties
        
        core_properties.title = act_type
        core_properties.author = "Cabinet STERU BARATTE AARPI"
        core_properties.subject = f"Dossier {context.get('case_number', 'XXX')}"
        core_properties.keywords = "droit pénal, défense, avocat"
        core_properties.category = "Acte juridique"
        core_properties.comments = f"Généré par Assistant Pénal le {datetime.now().strftime('%d/%m/%Y')}"
    
    def _process_communicated_pieces(self) -> List[Dict[str, Any]]:
        """Copie et renomme les pièces communiquées."""
        pieces_info = []
        
        for num, piece in self.pieces_registry.items():
            if piece.type != "versee":
                continue
            
            # Construire le nouveau nom
            # Format : 001_Titre_piece.ext
            title_clean = re.sub(r'[^\w\s-]', '', piece.titre)
            title_clean = re.sub(r'[-\s]+', '_', title_clean)
            title_clean = title_clean[:50]  # Limiter la longueur
            
            # Déterminer l'extension
            source_path = Path(piece.fichier_source)
            extension = source_path.suffix if source_path.suffix else '.pdf'
            
            new_filename = f"{num:03d}_{title_clean}{extension}"
            new_path = self.pieces_dir / new_filename
            
            # Copier le fichier si possible
            if source_path.exists():
                try:
                    shutil.copy2(source_path, new_path)
                    piece.fichier_communique = str(new_path)
                    
                    # Calculer le hash
                    with open(new_path, 'rb') as f:
                        piece.hash = hashlib.sha256(f.read()).hexdigest()
                    
                    pieces_info.append({
                        'numero': num,
                        'titre': piece.titre,
                        'fichier_original': piece.fichier_source,
                        'fichier_communique': str(new_path),
                        'hash': piece.hash
                    })
                    
                except Exception as e:
                    print(f"Erreur copie pièce {num} : {e}")
            
        return pieces_info
    
    def _summarize_document(self, chunks: List[Dict[str, Any]]) -> str:
        """Résume un document à partir de ses chunks."""
        # Concaténer les premiers chunks
        text = " ".join([chunk['content'] for chunk in chunks[:5]])
        
        # Résumé simple pour la démo
        if len(text) > 500:
            return text[:500] + "..."
        
        return text
    
    def _extract_key_facts(self, chunks: List[Dict[str, Any]]) -> List[str]:
        """Extrait les faits clés des chunks."""
        facts = []
        
        # Patterns pour identifier les faits importants
        fact_patterns = [
            r'[Ii]l résulte[^\.]+\.',
            r'[Ii]l est établi[^\.]+\.',
            r'[Ss]elon[^,]+,[^\.]+\.',
            r'[Ll]e \d+[^,]+,[^\.]+\.',  # Dates
        ]
        
        for chunk in chunks[:10]:
            text = chunk['content']
            
            for pattern in fact_patterns:
                matches = re.findall(pattern, text)
                facts.extend(matches)
        
        # Dédupliquer et limiter
        unique_facts = list(dict.fromkeys(facts))
        
        return unique_facts[:10]
    
    def _count_words(self, structured_content: Dict[str, Any]) -> int:
        """Compte les mots dans le contenu structuré."""
        total_words = 0
        
        for section in structured_content.get('sections', []):
            content = section.get('content', '')
            words = content.split()
            total_words += len(words)
        
        return total_words
    
    def _estimate_pages(self, structured_content: Dict[str, Any]) -> int:
        """Estime le nombre de pages du document."""
        # Estimation : ~250 mots par page
        word_count = self._count_words(structured_content)
        
        return max(1, word_count // 250)


# Export
__all__ = ['LegalDocumentGenerator', 'Piece', 'Citation']
